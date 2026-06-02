from __future__ import annotations

import csv
import html as html_module
import json
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from io import BytesIO, StringIO
from typing import Any

from app.core.heading_match import normalize_heading_match_text

try:
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable, _Cell as DocxCell  # type: ignore[attr-defined]
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.oxml.ns import qn as docx_qn  # type: ignore[attr-defined]
except ImportError:  # pragma: no cover
    DocxDocument = None  # type: ignore
    DocxTable = None  # type: ignore
    DocxCell = None  # type: ignore
    DocxParagraph = None  # type: ignore
    docx_qn = None  # type: ignore

try:
    from openpyxl import load_workbook
except ImportError:  # pragma: no cover
    load_workbook = None  # type: ignore

try:
    import xlrd
except ImportError:  # pragma: no cover
    xlrd = None  # type: ignore

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None  # type: ignore


class UnsupportedDocumentError(Exception):
    """无法解析或格式暂不支持（需在业务层转统一错误）。"""

    def __init__(self, message: str) -> None:
        super().__init__(message)
        self.message = message


@dataclass
class ParsedSegment:
    text: str
    start_offset: int
    end_offset: int
    page_no: int | None = None
    heading_path: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    parser_type: str
    text: str
    char_count: int
    segments: list[ParsedSegment]
    metadata: dict[str, Any] = field(default_factory=dict)


# 索引前安全阈值：兜底解析易把 PDF 流里的字面量拼成超长乱码
MAX_INDEXABLE_DOCUMENT_CHARS = 3_000_000
PDF_FALLBACK_SUSPICIOUS_CHARS = 80_000


def _cjk_ratio(sample: str) -> float:
    if not sample:
        return 0.0
    cjk = sum(1 for ch in sample if "\u4e00" <= ch <= "\u9fff")
    return cjk / len(sample)


def validate_parsed_document_for_indexing(parsed: ParsedDocument) -> None:
    """拒绝明显无效的解析结果，避免对乱码做上万次分块与向量化。"""
    text = (parsed.text or "").strip()
    if not text:
        return

    length = len(text)
    meta = parsed.metadata or {}
    backend = str(meta.get("parser_backend") or "")

    if length > MAX_INDEXABLE_DOCUMENT_CHARS:
        raise UnsupportedDocumentError(
            f"解析文本过长（约 {length:,} 字符），疑似误提取了 PDF 内部数据而非正文。"
            "请在 .env 中设置 MINERU_ENABLED=true 并启动 mineru 服务后重新解析，"
            "或提供可选中文本的 PDF。"
        )

    sample = text[: min(length, 80_000)]
    cjk_ratio = _cjk_ratio(sample)

    if backend == "pypdf_fallback":
        if length > PDF_FALLBACK_SUSPICIOUS_CHARS and cjk_ratio < 0.08:
            raise UnsupportedDocumentError(
                "PDF 通过兜底方式提取的正文疑似乱码（中文占比过低且文本极长）。"
                "该文件可能为扫描件或复杂版式，请在 .env 中设置 MINERU_ENABLED=true "
                "并启动 mineru 服务后重新解析。"
            )
        if length > 30_000 and len(parsed.segments) <= 1 and cjk_ratio < 0.12:
            raise UnsupportedDocumentError(
                "PDF 兜底解析得到单段超长文本且中文占比偏低，无法可靠分块。"
                "请启用 MinerU 解析（MINERU_ENABLED=true）后重试。"
            )
        if length > 20_000 and cjk_ratio < 0.02:
            raise UnsupportedDocumentError(
                "PDF 兜底解析未得到有效中文正文，请启用 MinerU 后重新解析。"
            )

    if parsed.parser_type == "pdf" and length > 500_000 and cjk_ratio < 0.05:
        raise UnsupportedDocumentError(
            "PDF 解析文本过长且中文占比异常，可能未正确提取正文。"
            "请启用 MinerU（MINERU_ENABLED=true）后重新解析。"
        )


def _decode_text_content(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _split_markdown_sections(text: str) -> list[tuple[str, str | None]]:
    lines = text.splitlines()
    sections: list[tuple[str, str | None]] = []
    heading_stack: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        content = "\n".join(buffer).strip()
        if content:
            sections.append((content, " / ".join(heading_stack) or None))
        buffer.clear()

    for line in lines:
        match = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
        if match:
            flush()
            level = len(match.group(1))
            heading_text = match.group(2).strip()
            heading_stack[:] = heading_stack[: level - 1]
            heading_stack.append(heading_text)
            buffer.append(line)
        else:
            buffer.append(line)

    flush()
    if not sections and text.strip():
        sections.append((text.strip(), None))
    return sections


def _build_segments_from_sections(sections: list[tuple[str, str | None]]) -> list[ParsedSegment]:
    segments: list[ParsedSegment] = []
    offset = 0
    for section_text, heading_path in sections:
        text = section_text.strip()
        if not text:
            continue
        start_offset = offset
        end_offset = start_offset + len(text)
        segments.append(
            ParsedSegment(
                text=text,
                start_offset=start_offset,
                end_offset=end_offset,
                heading_path=heading_path,
                metadata={"heading_path": heading_path},
            )
        )
        offset = end_offset + 2
    return segments


def _decode_pdf_escape_sequences(value: str) -> str:
    def _octal_replace(match: re.Match[str]) -> str:
        return chr(int(match.group(1), 8))

    value = re.sub(r"\\([0-7]{1,3})", _octal_replace, value)
    value = (
        value.replace(r"\\n", "\n")
        .replace(r"\\r", "\r")
        .replace(r"\\t", "\t")
        .replace(r"\\b", "\b")
        .replace(r"\\f", "\f")
        .replace(r"\\(", "(")
        .replace(r"\\)", ")")
        .replace(r"\\\\", "\\")
    )
    return value


def _extract_pdf_text_fallback(file_bytes: bytes) -> str:
    stream_pattern = re.compile(rb"stream\r?\n(.*?)\r?\nendstream", re.DOTALL)
    literal_pattern = re.compile(r"\((.*?)(?<!\\)\)")
    texts: list[str] = []

    for stream_match in stream_pattern.finditer(file_bytes):
        stream_bytes = stream_match.group(1)
        candidates = [stream_bytes]
        try:
            import zlib

            candidates.append(zlib.decompress(stream_bytes))
        except Exception:
            pass

        for candidate in candidates:
            decoded = candidate.decode("latin-1", errors="ignore")
            if "Tj" not in decoded and "TJ" not in decoded:
                continue
            for raw in literal_pattern.findall(decoded):
                text = _decode_pdf_escape_sequences(raw).strip()
                if text:
                    texts.append(text)

    merged = "\n".join(texts)
    merged = re.sub(r"\s+", " ", merged).strip()
    return merged


def _pdf_segment_substantive(segment: ParsedSegment) -> bool:
    """判断分段是否含可索引正文（排除纯图片占位等）。"""
    text = (segment.text or "").strip()
    if len(text) < 12:
        return False
    meta = segment.metadata or {}
    if meta.get("block") == "image":
        compact = re.sub(r"\s+", "", text)
        if compact in {"![Image]()", "![Image]", "[图片]"}:
            return False
        if "![Image]" in text and len(text) < 48:
            return False
    return True


def _pdf_text_mostly_image_markdown(text: str) -> bool:
    """ODL 扫描件常见：Markdown 仅图片占位，无可索引正文。"""
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return True
    image_like = 0
    for ln in lines:
        if re.match(r"^!\[[^\]]*\]\([^)]*\)\s*$", ln):
            image_like += 1
        elif "![Image]" in ln or ln in {"[图片]", "[Image]"}:
            image_like += 1
    return image_like >= max(1, int(len(lines) * 0.75))


def _pdf_parse_usable(doc: ParsedDocument) -> bool:
    """解析结果是否值得入库（避免 ODL 仅图片占位却 char_count>0 的情况）。"""
    text = (doc.text or "").strip()
    if not text:
        return False
    meta = doc.metadata or {}
    backend = str(meta.get("parser_backend") or "")
    segments = doc.segments or []

    if backend in {"opendataloader", "mineru"}:
        if not segments:
            return False
        if backend == "opendataloader" and int(meta.get("content_list_length") or 0) == 0:
            return False

    if not segments:
        if _pdf_text_mostly_image_markdown(text):
            return False
        return len(text) >= 32

    substantive = sum(1 for seg in segments if _pdf_segment_substantive(seg))
    if substantive > 0:
        return True
    if backend in {"pypdf", "pypdf_fallback"} and len(text) >= 64:
        return True
    if _pdf_text_mostly_image_markdown(text):
        return False
    return False


def _try_parse_pdf_opendataloader(
    file_bytes: bytes,
    fname: str,
    *,
    pdf_parser_hybrid: bool | None,
    log: Callable[[str], None] | None,
) -> ParsedDocument | None:
    from app.core.opendataloader_gateway import (
        OpenDataLoaderUnavailableError,
        get_opendataloader_gateway,
    )

    odl = get_opendataloader_gateway()
    if not odl.is_enabled():
        if log:
            log("OpenDataLoader 未启用，跳过")
        return None
    try:
        if log:
            log("使用 OpenDataLoader 解析 PDF…")
        result = odl.parse(
            file_bytes,
            fname,
            use_hybrid=pdf_parser_hybrid,
            log=log,
        )
        if log:
            log("OpenDataLoader：正在映射为内部分段结构…")
        doc = result.to_parsed_document()
        m = dict(doc.metadata or {})
        m["_parse_markdown"] = result.markdown if result.markdown is not None else ""
        m["_parse_content_list"] = result.to_content_list()
        doc.metadata = m
        if log:
            log(
                f"OpenDataLoader：字符数 {doc.char_count}，分段 {len(doc.segments)}，"
                f"标题 {doc.metadata.get('heading_count')}，表格 {doc.metadata.get('table_count')}"
            )
        return doc
    except OpenDataLoaderUnavailableError as e:
        if log:
            log(f"OpenDataLoader 不可用：{e}")
        return None


def _try_parse_pdf_mineru(
    file_bytes: bytes,
    fname: str,
    *,
    log: Callable[[str], None] | None,
) -> ParsedDocument | None:
    try:
        from app.core.mineru_gateway import MineruUnavailableError, get_mineru_gateway
    except Exception:  # pragma: no cover
        if log:
            log("MinerU 依赖不可用，跳过")
        return None

    gw = get_mineru_gateway()
    if not gw.is_enabled():
        if log:
            log("MinerU 未启用，跳过")
        return None
    if not gw.should_handle("pdf"):
        if log:
            log("MinerU 未配置接管 PDF，跳过")
        return None
    try:
        if log:
            log(f"使用 MinerU 解析 PDF（backend={gw.backend}, lang={gw.lang}）…")
        result = gw.parse(file_bytes, fname, log=log)
        if log:
            log("MinerU：正在映射为内部分段结构（标题路径、表格等）…")
        doc = result.to_parsed_document()
        m = dict(doc.metadata or {})
        m["_parse_markdown"] = result.markdown if result.markdown is not None else ""
        m["_parse_content_list"] = list(result.content_list)
        doc.metadata = m
        if log:
            log(
                f"MinerU：字符数 {doc.char_count}，分段 {len(doc.segments)}，"
                f"标题 {doc.metadata.get('heading_count')}，表格 {doc.metadata.get('table_count')}"
            )
        return doc
    except MineruUnavailableError as e:
        if log:
            log(f"MinerU 不可用：{e}")
        return None


def _try_parse_pdf_pypdf(file_bytes: bytes, *, log: Callable[[str], None] | None) -> ParsedDocument | None:
    if log:
        log("使用 pypdf 解析 PDF…")
    if PdfReader is not None:
        try:
            reader = PdfReader(BytesIO(file_bytes))
            segments: list[ParsedSegment] = []
            texts: list[str] = []
            offset = 0

            for index, page in enumerate(reader.pages, start=1):
                page_text = (page.extract_text() or "").strip()
                if not page_text:
                    continue
                texts.append(page_text)
                start_offset = offset
                end_offset = start_offset + len(page_text)
                segments.append(
                    ParsedSegment(
                        text=page_text,
                        start_offset=start_offset,
                        end_offset=end_offset,
                        page_no=index,
                        metadata={"page_no": index},
                    )
                )
                offset = end_offset + 2

            full_text = "\n\n".join(texts).strip()
            if full_text:
                return ParsedDocument(
                    parser_type="pdf",
                    text=full_text,
                    char_count=len(full_text),
                    segments=segments,
                    metadata={
                        "parser_backend": "pypdf",
                        "page_count": len(reader.pages),
                        "parsed_page_count": len(segments),
                    },
                )
        except Exception:
            pass

    fallback_text = _extract_pdf_text_fallback(file_bytes)
    if not fallback_text:
        return ParsedDocument(
            parser_type="pdf",
            text="",
            char_count=0,
            segments=[],
            metadata={"parser_backend": "pypdf_fallback", "fallback": True},
        )

    segment = ParsedSegment(
        text=fallback_text,
        start_offset=0,
        end_offset=len(fallback_text),
        metadata={"fallback": True},
    )
    return ParsedDocument(
        parser_type="pdf",
        text=fallback_text,
        char_count=len(fallback_text),
        segments=[segment],
        metadata={"parser_backend": "pypdf_fallback", "fallback": True},
    )


def _parse_pdf(
    file_bytes: bytes,
    *,
    file_name: str | None = None,
    pdf_parser: str | None = None,
    pdf_parser_hybrid: bool | None = None,
    log: Callable[[str], None] | None = None,
) -> ParsedDocument:
    from app.core.parser_config import resolve_pdf_fallback_chain

    fname = file_name or "doc.pdf"
    chain = resolve_pdf_fallback_chain(pdf_parser)
    primary = chain[0] if chain else "opendataloader"
    last_doc: ParsedDocument | None = None

    for index, engine in enumerate(chain):
        if index > 0 and log:
            log(f"PDF 解析降级：尝试 {engine}（首选 {primary}）…")

        doc: ParsedDocument | None = None
        if engine == "opendataloader":
            doc = _try_parse_pdf_opendataloader(
                file_bytes,
                fname,
                pdf_parser_hybrid=pdf_parser_hybrid,
                log=log,
            )
        elif engine == "mineru":
            doc = _try_parse_pdf_mineru(file_bytes, fname, log=log)
        elif engine == "pypdf":
            doc = _try_parse_pdf_pypdf(file_bytes, log=log)
        else:
            if log:
                log(f"PDF 引擎 {engine} 尚未实现，跳过")
            continue

        if doc is None:
            continue
        last_doc = doc

        if _pdf_parse_usable(doc):
            if index > 0:
                meta = dict(doc.metadata or {})
                meta["parser_fallback"] = True
                meta["parser_primary"] = primary
                meta["parser_backend_used"] = engine
                doc.metadata = meta
                if log:
                    log(f"PDF 已降级为 {engine} 解析成功")
            return doc

        if log:
            log(
                f"{engine} 解析结果不可用（字符 {doc.char_count}，分段 {len(doc.segments)}），"
                "尝试下一解析器…"
            )

    if last_doc is not None:
        return last_doc
    return ParsedDocument(
        parser_type="pdf",
        text="",
        char_count=0,
        segments=[],
        metadata={"parser_backend": "pypdf_fallback", "fallback": True},
    )


# MinerU 接管的图片后缀；没有本地降级路径（没有轻量级 OCR 备份），未启用 MinerU 时拒收
_IMAGE_SUFFIXES = {"png", "jpg", "jpeg", "bmp", "tif", "tiff", "webp"}


def _parse_with_mineru(
    file_bytes: bytes,
    *,
    suffix: str,
    file_name: str | None = None,
    parser_type: str,
    log: Callable[[str], None] | None = None,
    extra_metadata: dict[str, Any] | None = None,
    require_format_whitelist: bool = True,
) -> ParsedDocument:
    """
    通过 MinerU /file_parse 解析并在 metadata 中附带 Markdown / content_list 侧车字段。
    未启用、后缀不在白名单（仅 require_format_whitelist 时）或服务不可达时抛 UnsupportedDocumentError。
    """
    try:
        from app.core.mineru_gateway import MineruUnavailableError, get_mineru_gateway
    except Exception as e:  # pragma: no cover
        raise UnsupportedDocumentError(f"MinerU 依赖不可用：{e}") from e

    gw = get_mineru_gateway()
    if not gw.is_enabled():
        raise UnsupportedDocumentError(
            "未启用 MinerU。请在部署中打开 MINERU_ENABLED 并拉起 mineru 服务。"
        )
    if require_format_whitelist and not gw.should_handle(suffix):
        raise UnsupportedDocumentError(
            f"MinerU 未配置接管 .{suffix} 后缀。请在 MINERU_FORMATS 中加入。"
        )

    fname = file_name or f"document.{suffix}"
    if log:
        log(f"使用 MinerU 解析（.{suffix}, backend={gw.backend}, lang={gw.lang}）…")
    try:
        result = gw.parse(file_bytes, fname, log=log)
    except MineruUnavailableError as e:
        raise UnsupportedDocumentError(f"MinerU 不可达，解析失败：{e}") from e

    if log:
        log("MinerU：正在映射为内部分段结构（标题路径、表格等）…")
    doc = result.to_parsed_document()
    doc.parser_type = parser_type
    meta = dict(doc.metadata or {})
    meta["parser_backend"] = "mineru"
    if extra_metadata:
        meta.update(extra_metadata)
    meta["_parse_markdown"] = result.markdown if result.markdown is not None else ""
    meta["_parse_content_list"] = list(result.content_list)
    doc.metadata = meta
    if log:
        log(
            f"MinerU（{parser_type}）：字符数 {doc.char_count}，分段 {len(doc.segments)}，"
            f"标题 {meta.get('heading_count')}，表格 {meta.get('table_count')}"
        )
    return doc


def _parse_image(
    file_bytes: bytes,
    *,
    suffix: str,
    file_name: str | None = None,
    log: Callable[[str], None] | None = None,
) -> ParsedDocument:
    """图片解析完全依赖 MinerU OCR + 版面分析，无本地降级路径。"""
    return _parse_with_mineru(
        file_bytes,
        suffix=suffix,
        file_name=file_name,
        parser_type=f"image.{suffix}",
        log=log,
        extra_metadata={"image_suffix": suffix},
    )


# 章节标题识别（启发式）：优先依赖 Word Heading 样式；无样式回退到中文/数字特征。
# 用"等级类（class）"而非绝对级别，避免章/节/条混杂时栈层级错乱。
_CN_NUM = r"[一二三四五六七八九十百千]+"
# 启发式标题：不含 item 级（「1、xxx」列表项），避免办事指南枚举被当成标题
_HEADING_CLASSES: list[tuple[str, "re.Pattern[str]"]] = [
    # chapter（章）：第X章、一、背景介绍、第1章
    ("chapter", re.compile(rf"^第\s*{_CN_NUM}\s*[章节篇部].{{0,60}}$")),
    ("chapter", re.compile(rf"^{_CN_NUM}\s*[、.．].{{0,60}}$")),
    ("chapter", re.compile(r"^第\s*\d+\s*[章节篇部条款].{0,60}$")),
    # section（节）：（一）xxx、1.1 xxx、1.2.3 xxx
    ("section", re.compile(rf"^[(（]\s*{_CN_NUM}\s*[)）].{{0,60}}$")),
    ("section", re.compile(r"^\d+(?:\.\d+){1,3}\s+.{0,60}$")),
]

# 列表项形如「1、xxx」「2. xxx」——不当标题，仅保留 Word Heading 3+ 样式识别 item
_ENUM_LIST_LINE = re.compile(r"^\d+\s*[、.．)）]\s*")

_CLASS_ORDER = {"chapter": 1, "section": 2, "item": 3}


def _style_heading_class(style_name: str | None) -> str | None:
    if not style_name:
        return None
    m = re.match(r"^(?:Heading|标题)\s*(\d+)$", style_name.strip(), re.IGNORECASE)
    if not m:
        return None
    level = int(m.group(1))
    if level <= 1:
        return "chapter"
    if level == 2:
        return "section"
    if level >= 3:
        return "item"
    return None


def _detect_heading_class(text: str, style_name: str | None) -> str | None:
    cls = _style_heading_class(style_name)
    if cls is not None:
        return cls
    t = text.strip()
    if not t or len(t) > 60:
        return None
    t_norm = normalize_heading_match_text(text)
    if _ENUM_LIST_LINE.match(t_norm):
        return None
    for class_tag, regex in _HEADING_CLASSES:
        if regex.match(t_norm):
            return class_tag
    return None


def _update_heading_path(stack: list[tuple[str, str]], heading_class: str, text: str) -> str:
    """
    按 class 等级维护标题栈：
    - 同 class → 同级替换；
    - 更低 class → 保留；
    - 更高或同级 → 截断。
    """
    cur_order = _CLASS_ORDER.get(heading_class, 3)
    new_stack: list[tuple[str, str]] = []
    for cls, t in stack:
        if _CLASS_ORDER.get(cls, 3) < cur_order:
            new_stack.append((cls, t))
        else:
            break
    new_stack.append((heading_class, text.strip()))
    stack[:] = new_stack
    return " / ".join(t for _, t in new_stack)


def _docx_ancestor_skips_content(el: Any, paragraph_el: Any) -> bool:
    """修订 / 移动删除区内的节点不纳入终稿文本（与 Word 修订「最终状态」常见展示一致）。"""
    if docx_qn is None:
        return False
    del_tag = docx_qn("w:del")
    move_from_tag = docx_qn("w:moveFrom")
    parent = el.getparent()
    while parent is not None:
        if parent == paragraph_el:
            return False
        if parent.tag in (del_tag, move_from_tag):
            return True
        parent = parent.getparent()
    return False


def _docx_paragraph_plain_text(paragraph: Any) -> str:
    """
    从段落 w:p 的 OOXML 子树抽取可见文本：
    汇总 w:t、w:tab、w:br（及 w:cr），覆盖超链接/内容控件等路径下 python-docx 的 paragraph.text 易漏字的情况。
    """
    if docx_qn is None:
        return str(getattr(paragraph, "text", "") or "")
    p_el = paragraph._element  # type: ignore[attr-defined]
    w_t = docx_qn("w:t")
    w_tab = docx_qn("w:tab")
    w_br = docx_qn("w:br")
    w_cr = docx_qn("w:cr")
    parts: list[str] = []
    for el in p_el.iter():
        if el.tag == w_t:
            if _docx_ancestor_skips_content(el, p_el):
                continue
            if el.text:
                parts.append(el.text)
        elif el.tag == w_tab:
            if _docx_ancestor_skips_content(el, p_el):
                continue
            parts.append("\t")
        elif el.tag in (w_br, w_cr):
            if _docx_ancestor_skips_content(el, p_el):
                continue
            parts.append("\n")
    return "".join(parts)


def _docx_table_cell_text(cell: "DocxCell") -> str:
    """把单元格内所有段落拼成一行文本（去换行，保留空格）；抽字走 OOXML 与正文段落一致。"""
    try:
        paragraphs = cell.paragraphs  # type: ignore[attr-defined]
    except Exception:
        return (cell.text or "").strip().replace("\n", " ")
    pieces: list[str] = []
    for p in paragraphs:
        raw = _docx_paragraph_plain_text(p).strip()
        if raw:
            pieces.append(raw)
    if pieces:
        return " ".join(pieces)
    return (cell.text or "").strip().replace("\n", " ")


def _docx_page_break_count_in_element(el: Any) -> int:
    """统计 OOXML 子树中的 Word 分页标记（lastRenderedPageBreak + 硬分页 w:br type=page）。"""
    if docx_qn is None:
        return 0
    lrpb = docx_qn("w:lastRenderedPageBreak")
    br_tag = docx_qn("w:br")
    br_type = docx_qn("w:type")
    count = 0
    for child in el.iter():
        if child is el:
            continue
        if child.tag == lrpb:
            count += 1
        elif child.tag == br_tag and child.get(br_type) == "page":
            count += 1
    return count


def _docx_leading_page_breaks(paragraph: Any) -> int:
    """段落第一个可见文字之前的分页符数量（Word 软分页常出现在段首）。"""
    if docx_qn is None:
        try:
            return len(paragraph.rendered_page_breaks)
        except Exception:
            return 0
    p_el = paragraph._element  # type: ignore[attr-defined]
    lrpb = docx_qn("w:lastRenderedPageBreak")
    br_tag = docx_qn("w:br")
    br_type = docx_qn("w:type")
    w_t = docx_qn("w:t")
    count = 0
    for child in p_el.iter():
        if child is p_el:
            continue
        if child.tag == w_t and (child.text or "").strip():
            break
        if child.tag == lrpb:
            count += 1
        elif child.tag == br_tag and child.get(br_type) == "page":
            count += 1
    return count


def _docx_apply_page_breaks_before_block(current_page: int, node: Any, *, kind: str) -> int:
    """处理仅含分页符的空段落，或表格/段落前的分页增量。"""
    if kind == "paragraph":
        text = _docx_paragraph_plain_text(node).strip() or str(getattr(node, "text", "") or "").strip()
        if text:
            return current_page
        return current_page + _docx_page_break_count_in_element(node._element)  # type: ignore[attr-defined]
    return current_page


def _docx_finalize_paragraph_page(current_page: int, paragraph: Any) -> tuple[int, int]:
    """返回 (segment_page_no, next_current_page)。"""
    leading = _docx_leading_page_breaks(paragraph)
    page_no = current_page + leading
    total_breaks = _docx_page_break_count_in_element(paragraph._element)  # type: ignore[attr-defined]
    trailing = max(0, total_breaks - leading)
    return page_no, page_no + trailing


def _docx_finalize_table_page(current_page: int, table: Any) -> tuple[int, int]:
    page_no = current_page
    breaks = _docx_page_break_count_in_element(table._element)  # type: ignore[attr-defined]
    return page_no, page_no + breaks


def _is_plausible_header_row(cells: list[str]) -> bool:
    """启发式：若首行均不为空、且没有显著长于其他行的值，视为表头。"""
    if not cells or any(not c for c in cells):
        return False
    if any(len(c) > 40 for c in cells):
        return False
    return True


MAX_WHOLE_TABLE_TEXT_CHARS = 12_000
MAX_WHOLE_TABLE_DATA_ROWS = 100
# 宽表按行检索：单块字符预算 + 每块最多行数；段数硬顶避免索引爆炸
MAX_TABLE_ROW_SEGMENT_CHARS = 4_000
MAX_TABLE_ROWS_PER_SEGMENT = 10
MAX_TABLE_ROW_SEGMENTS = 2_400


def _rows_to_markdown_table(rows: list[list[str]]) -> str:
    cleaned = [[c.strip() for c in row] for row in rows if any(c.strip() for c in row)]
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    norm = [r + [""] * (width - len(r)) for r in cleaned]
    lines = ["| " + " | ".join(r) + " |" for r in norm]
    if len(norm) >= 2:
        sep = "| " + " | ".join("---" for _ in range(width)) + " |"
        lines.insert(1, sep)
    return "\n".join(lines)


def _table_body_retrieval_text(body: str, rows: list[list[str]]) -> str:
    """MinerU / ODL 表格侧车正文 → 索引用纯文本（优先保留 Markdown 表格）。"""
    stripped = body.strip()
    if stripped.startswith("|") or (stripped.startswith("<") is False and "| ---" in stripped):
        return stripped
    if stripped.startswith("<"):
        md = _rows_to_markdown_table(rows)
        return md or stripped
    md = _rows_to_markdown_table(rows)
    return md or stripped


def prepend_table_context(retrieval_text: str, context_prefix: str | None) -> str:
    """将同页标题/附件等上下文前缀并入表格检索文本，避免表格块缺少语义。"""
    body = (retrieval_text or "").strip()
    prefix = (context_prefix or "").strip()
    if not prefix:
        return body
    if not body:
        return prefix
    if body.startswith(prefix) or prefix in body.split("\n\n", 1)[0]:
        return body
    probe = prefix[: min(24, len(prefix))]
    if probe and probe in body:
        return body
    return f"{prefix}\n\n{body}"


def build_table_segments_smart(
    rows: list[list[str]],
    ti: int,
    offset_ref: list[int],
    parts: list[str],
    heading_path: str | None,
    *,
    table_body: str | None = None,
    table_body_html: str | None = None,
    page_no: int | None = None,
    context_prefix: str | None = None,
) -> list[ParsedSegment]:
    """
    表格分段：中等规模整表一块（与 Markdown 预览一致）；超大表按批合并行，避免单行一块。
    """
    normalized = [r for r in rows if any(c.strip() for c in r)]
    if not normalized:
        return []

    body = (table_body or "").strip()
    html = (table_body_html or "").strip()
    if not html and body.startswith("<"):
        html = body
    if not html and normalized:
        html = _rows_to_html_table(normalized)

    data_row_count = _table_data_row_count(normalized)
    retrieval_text = prepend_table_context(_table_body_retrieval_text(body, normalized), context_prefix)

    if (
        retrieval_text
        and len(retrieval_text) <= MAX_WHOLE_TABLE_TEXT_CHARS
        and data_row_count <= MAX_WHOLE_TABLE_DATA_ROWS
    ):
        start = offset_ref[0]
        end = start + len(retrieval_text)
        meta: dict[str, Any] = {
            "block": "table",
            "table_index": ti,
            "table_role": "body",
            "table_header": normalized[0] if len(normalized) >= 2 and _is_plausible_header_row(normalized[0]) else None,
            "table_row_count": data_row_count,
            "heading_path": heading_path,
        }
        if html:
            meta["table_body_html"] = html
        if page_no is not None:
            meta["page_no"] = page_no
            meta["page_idx"] = page_no - 1
        if context_prefix and context_prefix.strip():
            meta["table_context_prefix"] = context_prefix.strip()
        seg = ParsedSegment(
            text=retrieval_text,
            start_offset=start,
            end_offset=end,
            page_no=page_no,
            heading_path=heading_path,
            metadata=meta,
        )
        parts.append(retrieval_text)
        offset_ref[0] = end + 1
        return [seg]

    row_batch = _table_row_batch_size(normalized)
    return build_table_segments_from_rows(
        normalized,
        ti,
        offset_ref,
        parts,
        heading_path,
        table_body_html=html or None,
        row_batch_size=row_batch,
        page_no=page_no,
        context_prefix=context_prefix,
    )


def build_table_segments_from_rows(
    rows: list[list[str]],
    ti: int,
    offset_ref: list[int],
    parts: list[str],
    heading_path: str | None,
    *,
    table_body_html: str | None = None,
    row_batch_size: int = 1,
    page_no: int | None = None,
    context_prefix: str | None = None,
) -> list[ParsedSegment]:
    """
    通用"行 → 表格 ParsedSegment 列表"转换器：
    - 第 1 个 segment：表格概览（表头 + 行数），便于"这个表格是关于什么"类召回
    - 每个数据行：一个 segment，文本为 "列1：值1；列2：值2"（若有表头）或 "值1 | 值2..."（无表头）

    offset_ref 是 [current_offset] 的可变引用，用于累加偏移，保持与全文 offset 一致。
    docx、MinerU PDF 表格等不同来源都应归一到 list[list[str]] 后调用本函数。
    """
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return []

    header: list[str] | None = None
    if len(rows) >= 2 and _is_plausible_header_row(rows[0]):
        header = rows[0]
        data_rows = rows[1:]
    else:
        data_rows = rows

    segments: list[ParsedSegment] = []

    overview_parts: list[str] = [f"表格 {ti + 1}"]
    if header:
        overview_parts.append("列：" + " | ".join(header))
    overview_parts.append(f"共 {len(data_rows)} 行")
    overview = "；".join(overview_parts)
    start = offset_ref[0]
    end = start + len(overview)
    segments.append(
        ParsedSegment(
            text=overview,
            start_offset=start,
            end_offset=end,
            page_no=page_no,
            heading_path=heading_path,
            metadata={
                "block": "table",
                "table_index": ti,
                "table_role": "overview",
                "table_header": header,
                "table_row_count": len(data_rows),
                "heading_path": heading_path,
                **({"page_no": page_no, "page_idx": page_no - 1} if page_no is not None else {}),
            },
        )
    )
    parts.append(overview)
    offset_ref[0] = end + 1

    batch_size = max(1, int(row_batch_size))
    prefix = (context_prefix or "").strip()
    prefix_applied = False

    ri = 0
    while ri < len(data_rows):
        batch = data_rows[ri : ri + batch_size]
        lines = [_format_table_row_line(cells, header) for cells in batch]
        lines = [ln for ln in lines if ln]
        if not lines:
            ri += batch_size
            continue
        line = "\n".join(lines)
        row_end = ri + len(batch) - 1
        start = offset_ref[0]
        end = start + len(line)
        html_rows = ([header] + batch) if header else batch
        batch_html = _rows_to_html_table(html_rows) if html_rows else None
        row_meta: dict[str, Any] = {
            "block": "table",
            "table_index": ti,
            "table_role": "row",
            "table_row_index": ri,
            "table_row_index_end": row_end,
            "table_row_batch_size": len(batch),
            "table_header": header,
            "heading_path": heading_path,
        }
        if batch_html:
            row_meta["table_body_html"] = batch_html
        if page_no is not None:
            row_meta["page_no"] = page_no
            row_meta["page_idx"] = page_no - 1
        if prefix and not prefix_applied:
            line = prepend_table_context(line, prefix)
            row_meta["table_context_prefix"] = prefix
            prefix_applied = True
            end = start + len(line)
        segments.append(
            ParsedSegment(
                text=line,
                start_offset=start,
                end_offset=end,
                page_no=page_no,
                heading_path=heading_path,
                metadata=row_meta,
            )
        )
        parts.append(line)
        offset_ref[0] = end + 1
        ri += batch_size

    return segments


def compact_table_segments_for_index(
    segments: list[ParsedSegment],
    *,
    max_segments: int = 2400,
) -> list[ParsedSegment]:
    """索引前压缩过密的表格行 segment（兼容旧解析结果或未重启服务的情况）。"""
    if len(segments) <= max_segments:
        return segments

    passthrough: list[ParsedSegment] = []
    row_by_table: dict[int, list[ParsedSegment]] = {}

    for seg in segments:
        meta = seg.metadata or {}
        if meta.get("block") == "table" and meta.get("table_role") == "row":
            ti = int(meta.get("table_index") or 0)
            row_by_table.setdefault(ti, []).append(seg)
        else:
            passthrough.append(seg)

    merged_rows: list[ParsedSegment] = []
    table_count = max(1, len(row_by_table))
    per_table_budget = max(80, (max_segments - len(passthrough)) // table_count)

    for _ti, row_segs in sorted(row_by_table.items()):
        row_segs.sort(key=lambda s: (s.metadata or {}).get("table_row_index", 0))
        n = len(row_segs)
        batch = max(1, (n + per_table_budget - 1) // per_table_budget)
        for i in range(0, n, batch):
            group = row_segs[i : i + batch]
            text = "\n".join(s.text for s in group if s.text.strip())
            if not text.strip():
                continue
            first, last = group[0], group[-1]
            meta = dict(first.metadata or {})
            fmeta, lmeta = first.metadata or {}, last.metadata or {}
            meta["table_row_index"] = fmeta.get("table_row_index", i)
            meta["table_row_index_end"] = lmeta.get("table_row_index", i + len(group) - 1)
            meta["table_row_batch_size"] = len(group)
            merged_rows.append(
                ParsedSegment(
                    text=text,
                    start_offset=first.start_offset,
                    end_offset=last.end_offset,
                    heading_path=first.heading_path,
                    page_no=first.page_no,
                    metadata=meta,
                )
            )

    combined = passthrough + merged_rows
    combined.sort(key=lambda s: s.start_offset if s.start_offset is not None else 0)
    return combined


def _docx_table_to_segments(
    table: "DocxTable",
    ti: int,
    offset_ref: list[int],
    parts: list[str],
    heading_path: str | None,
    *,
    page_no: int | None = None,
) -> list[ParsedSegment]:
    """从 python-docx 的 Table 提取行，再交给通用的 build_table_segments_from_rows 处理。"""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_docx_table_cell_text(c) for c in row.cells]
        if any(cells):
            rows.append(cells)
    row_batch = _table_row_batch_size(rows) if len(rows) > 2 else 1
    return build_table_segments_from_rows(
        rows,
        ti,
        offset_ref,
        parts,
        heading_path,
        row_batch_size=row_batch,
        page_no=page_no,
    )


def _iter_docx_body_blocks(doc: Any):
    """按文档顺序迭代 body 中的段落与表格（保持原始 docx body 顺序）。"""
    if docx_qn is None or DocxParagraph is None or DocxTable is None:
        for p in doc.paragraphs:
            yield ("paragraph", p)
        for t in doc.tables:
            yield ("table", t)
        return
    body = doc.element.body
    p_tag = docx_qn("w:p")
    tbl_tag = docx_qn("w:tbl")
    for child in body.iterchildren():
        if child.tag == p_tag:
            yield ("paragraph", DocxParagraph(child, doc))
        elif child.tag == tbl_tag:
            yield ("table", DocxTable(child, doc))


def _parse_docx(file_bytes: bytes) -> ParsedDocument:
    if DocxDocument is None:
        raise UnsupportedDocumentError("未安装 python-docx，无法解析 .docx 文件")
    doc = DocxDocument(BytesIO(file_bytes))
    segments: list[ParsedSegment] = []
    parts: list[str] = []
    offset_ref = [0]
    heading_stack: list[tuple[str, str]] = []
    current_heading_path: str | None = None
    table_counter = 0
    current_page = 1

    for kind, node in _iter_docx_body_blocks(doc):
        if kind == "paragraph":
            t = _docx_paragraph_plain_text(node).strip()
            if not t:
                t = (node.text or "").strip()
            if not t:
                current_page = _docx_apply_page_breaks_before_block(current_page, node, kind=kind)
                continue
            page_no, current_page = _docx_finalize_paragraph_page(current_page, node)
            style_name = None
            try:
                style_name = node.style.name if node.style is not None else None  # type: ignore[attr-defined]
            except Exception:
                style_name = None
            cls = _detect_heading_class(t, style_name)
            start = offset_ref[0]
            end = start + len(t)
            page_meta = {"page_no": page_no, "page_idx": page_no - 1}
            if cls is not None:
                current_heading_path = _update_heading_path(heading_stack, cls, t)
                segments.append(
                    ParsedSegment(
                        text=t,
                        start_offset=start,
                        end_offset=end,
                        page_no=page_no,
                        heading_path=current_heading_path,
                        metadata={
                            "block": "heading",
                            "heading_class": cls,
                            "heading_level": _CLASS_ORDER.get(cls, 3),
                            "heading_path": current_heading_path,
                            **page_meta,
                        },
                    )
                )
            else:
                segments.append(
                    ParsedSegment(
                        text=t,
                        start_offset=start,
                        end_offset=end,
                        page_no=page_no,
                        heading_path=current_heading_path,
                        metadata={
                            "block": "paragraph",
                            "heading_path": current_heading_path,
                            **page_meta,
                        },
                    )
                )
            parts.append(t)
            offset_ref[0] = end + 1
        elif kind == "table":
            page_no, current_page = _docx_finalize_table_page(current_page, node)
            table_segments = _docx_table_to_segments(
                node,
                table_counter,
                offset_ref,
                parts,
                current_heading_path,
                page_no=page_no,
            )
            segments.extend(table_segments)
            table_counter += 1

    full_text = "\n".join(parts).strip()
    if not full_text:
        return ParsedDocument(parser_type="docx", text="", char_count=0, segments=[], metadata={})
    heading_count = sum(1 for s in segments if (s.metadata or {}).get("block") == "heading")
    table_count = table_counter
    table_row_count = sum(
        1 for s in segments if (s.metadata or {}).get("block") == "table" and (s.metadata or {}).get("table_role") == "row"
    )
    return ParsedDocument(
        parser_type="docx",
        text=full_text,
        char_count=len(full_text),
        segments=segments,
        metadata={
            "parser_backend": "python-docx",
            "paragraph_and_table_blocks": len(segments),
            "heading_count": heading_count,
            "table_count": table_count,
            "table_row_count": table_row_count,
        },
    )


def _parse_xlsx_xlsm(file_bytes: bytes, suffix: str, log: Callable[[str], None] | None = None) -> ParsedDocument:
    return _parse_excel(file_bytes, suffix, engine="tsv", log=log)


def _parse_xls(file_bytes: bytes, log: Callable[[str], None] | None = None) -> ParsedDocument:
    return _parse_excel(file_bytes, "xls", engine="tsv", log=log)


def _cell_str(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _rows_to_html_table(rows: list[list[str]], *, caption: str | None = None) -> str:
    if not rows:
        return ""
    parts = ["<table>"]
    if caption:
        parts.append(f"<caption>{html_module.escape(caption)}</caption>")
    for ri, row in enumerate(rows):
        parts.append("<tr>")
        tag = "th" if ri == 0 else "td"
        for cell in row:
            parts.append(f"<{tag}>{html_module.escape(cell)}</{tag}>")
        parts.append("</tr>")
    parts.append("</table>")
    return "".join(parts)


def _read_xlsx_sheet_rows(sheet, log: Callable[[str], None] | None = None) -> list[list[str]]:
    def _log(msg: str) -> None:
        if log:
            log(msg)

    rows: list[list[str]] = []
    row_count = 0
    for row in sheet.iter_rows(values_only=True):
        row_count += 1
        cells = [_cell_str(c) for c in row]
        if not any(cells):
            continue
        rows.append(cells)
        if row_count % 5000 == 0:
            _log(f"工作表「{sheet.title}」已扫描 {row_count} 行…")
    return rows


def _read_xls_sheet_rows(book, sheet_index: int, log: Callable[[str], None] | None = None) -> tuple[str, list[list[str]]]:
    def _log(msg: str) -> None:
        if log:
            log(msg)

    sheet = book.sheet_by_index(sheet_index)
    rows: list[list[str]] = []
    for ri in range(sheet.nrows):
        values = sheet.row_values(ri)
        cells = [_cell_str(c) for c in values]
        if not any(cells):
            continue
        rows.append(cells)
        if ri > 0 and ri % 5000 == 0:
            _log(f"工作表「{sheet.name}」已读取 {ri} 行…")
    return sheet.name, rows


def _format_table_row_line(cells: list[str], header: list[str] | None) -> str:
    if header and len(cells) == len(header):
        return "；".join(f"{h}：{v}" for h, v in zip(header, cells) if v).strip()
    return " | ".join(c for c in cells if c).strip()


def _split_table_header_and_data(rows: list[list[str]]) -> tuple[list[str] | None, list[list[str]]]:
    normalized = [r for r in rows if any(c.strip() for c in r)]
    if len(normalized) >= 2 and _is_plausible_header_row(normalized[0]):
        return normalized[0], normalized[1:]
    return None, normalized


def _estimate_table_row_chars(rows: list[list[str]], header: list[str] | None, *, sample_size: int = 48) -> int:
    """估算单行检索文本长度，用于按字符预算合并行 batch。"""
    _, data_rows = _split_table_header_and_data(rows)
    if not data_rows:
        return 200
    sample = data_rows[:sample_size]
    lengths = [
        len(_format_table_row_line(r, header))
        for r in sample
        if any(c.strip() for c in r) and _format_table_row_line(r, header)
    ]
    if not lengths:
        return 200
    return max(80, int(sum(lengths) / len(lengths)))


def _table_row_batch_size(
    rows: list[list[str]],
    *,
    max_segment_chars: int = MAX_TABLE_ROW_SEGMENT_CHARS,
    max_rows_per_batch: int = MAX_TABLE_ROWS_PER_SEGMENT,
    max_row_segments: int = MAX_TABLE_ROW_SEGMENTS,
) -> int:
    """
    宽表按字符预算切 batch，便于按行/目录名精确召回；
    段数超硬顶时再放大 batch，避免索引块数爆炸。
    """
    header, data_rows = _split_table_header_and_data(rows)
    num_data_rows = len(data_rows)
    if num_data_rows <= 1:
        return 1

    avg_chars = _estimate_table_row_chars(rows, header)
    by_chars = max(1, max_segment_chars // avg_chars)
    batch = min(by_chars, max_rows_per_batch)

    if num_data_rows > max_row_segments:
        min_batch_for_cap = (num_data_rows + max_row_segments - 1) // max_row_segments
        batch = max(batch, min_batch_for_cap)

    return max(1, batch)


def _table_data_row_count(row_matrix: list[list[str]]) -> int:
    rows = [r for r in row_matrix if any(c.strip() for c in r)]
    if len(rows) >= 2 and _is_plausible_header_row(rows[0]):
        return len(rows) - 1
    return len(rows)


def _parse_excel_tsv(file_bytes: bytes, suffix: str, log: Callable[[str], None] | None = None) -> ParsedDocument:
    if suffix in {"xlsx", "xlsm"}:
        if load_workbook is None:
            raise UnsupportedDocumentError("未安装 openpyxl，无法解析 Excel 文档")

        def _log(msg: str) -> None:
            if log:
                log(msg)

        _log(f"Excel（{suffix}，TSV）：openpyxl 只读，约 {len(file_bytes)} 字节")
        wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        try:
            segments: list[ParsedSegment] = []
            parts: list[str] = []
            offset = 0
            for sheet in wb:
                rows_text: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    cells = [_cell_str(c) for c in row]
                    if not any(cells):
                        continue
                    rows_text.append("\t".join(cells))
                sheet_body = "\n".join(rows_text).strip()
                if not sheet_body:
                    continue
                block = f"## {sheet.title}\n{sheet_body}"
                start = offset
                end = offset + len(block)
                segments.append(
                    ParsedSegment(
                        text=block,
                        start_offset=start,
                        end_offset=end,
                        heading_path=sheet.title,
                        metadata={"sheet": sheet.title, "excel_engine": "tsv"},
                    )
                )
                parts.append(block)
                offset = end + 2
            full_text = "\n\n".join(parts).strip()
            label = "xlsx" if suffix == "xlsx" else "xlsm"
            return ParsedDocument(
                parser_type=label,
                text=full_text,
                char_count=len(full_text),
                segments=segments
                or [ParsedSegment(text=full_text, start_offset=0, end_offset=len(full_text), metadata={})],
                metadata={"sheet_count": len(segments), "excel_engine": "tsv"},
            )
        finally:
            wb.close()

    if xlrd is None:
        raise UnsupportedDocumentError("未安装 xlrd，无法解析 .xls 文件")

    def _log_xls(msg: str) -> None:
        if log:
            log(msg)

    _log_xls(f"Excel（xls，TSV）：约 {len(file_bytes)} 字节")
    book = xlrd.open_workbook(file_contents=file_bytes)
    segments = []
    parts: list[str] = []
    offset = 0
    for si in range(book.nsheets):
        title, row_matrix = _read_xls_sheet_rows(book, si, log=log)
        rows_text = ["\t".join(r) for r in row_matrix]
        sheet_body = "\n".join(rows_text).strip()
        if not sheet_body:
            continue
        block = f"## {title}\n{sheet_body}"
        start = offset
        end = offset + len(block)
        segments.append(
            ParsedSegment(
                text=block,
                start_offset=start,
                end_offset=end,
                heading_path=title,
                metadata={"sheet": title, "excel_engine": "tsv"},
            )
        )
        parts.append(block)
        offset = end + 2
    full_text = "\n\n".join(parts).strip()
    return ParsedDocument(
        parser_type="xls",
        text=full_text,
        char_count=len(full_text),
        segments=segments or [ParsedSegment(text=full_text, start_offset=0, end_offset=len(full_text), metadata={})],
        metadata={"sheet_count": len(segments), "excel_engine": "tsv"},
    )


def _parse_excel_html_table(file_bytes: bytes, suffix: str, log: Callable[[str], None] | None = None) -> ParsedDocument:
    def _log(msg: str) -> None:
        if log:
            log(msg)

    segments: list[ParsedSegment] = []
    parts: list[str] = []
    offset_ref = [0]
    table_index = 0

    if suffix in {"xlsx", "xlsm"}:
        if load_workbook is None:
            raise UnsupportedDocumentError("未安装 openpyxl，无法解析 Excel 文档")
        _log(f"Excel（{suffix}，HTML 表格）：openpyxl 只读，约 {len(file_bytes)} 字节")
        wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
        try:
            for sheet in wb:
                row_matrix = _read_xlsx_sheet_rows(sheet, log=log)
                if not row_matrix:
                    _log(f"工作表「{sheet.title}」无有效内容，跳过")
                    continue
                table_html = _rows_to_html_table(row_matrix, caption=sheet.title)
                row_batch = _table_row_batch_size(row_matrix)
                if row_batch > 1 and log:
                    avg_chars = _estimate_table_row_chars(row_matrix, _split_table_header_and_data(row_matrix)[0])
                    _log(
                        f"工作表「{sheet.title}」约 {_table_data_row_count(row_matrix)} 行，"
                        f"均行 {avg_chars} 字，按每 {row_batch} 行合并为 1 个检索块（≤{MAX_TABLE_ROW_SEGMENT_CHARS} 字/块）"
                    )
                sheet_segments = build_table_segments_from_rows(
                    row_matrix,
                    table_index,
                    offset_ref,
                    parts,
                    sheet.title,
                    table_body_html=table_html,
                    row_batch_size=row_batch,
                )
                for seg in sheet_segments:
                    meta = dict(seg.metadata or {})
                    meta["sheet"] = sheet.title
                    meta["excel_engine"] = "html_table"
                    seg.metadata = meta
                segments.extend(sheet_segments)
                table_index += 1
        finally:
            wb.close()
    else:
        if xlrd is None:
            raise UnsupportedDocumentError("未安装 xlrd，无法解析 .xls 文件")
        _log(f"Excel（xls，HTML 表格）：约 {len(file_bytes)} 字节")
        book = xlrd.open_workbook(file_contents=file_bytes)
        for si in range(book.nsheets):
            title, row_matrix = _read_xls_sheet_rows(book, si, log=log)
            if not row_matrix:
                continue
            table_html = _rows_to_html_table(row_matrix, caption=title)
            row_batch = _table_row_batch_size(row_matrix)
            if row_batch > 1 and log:
                avg_chars = _estimate_table_row_chars(row_matrix, _split_table_header_and_data(row_matrix)[0])
                _log(
                    f"工作表「{title}」约 {_table_data_row_count(row_matrix)} 行，"
                    f"均行 {avg_chars} 字，按每 {row_batch} 行合并为 1 个检索块（≤{MAX_TABLE_ROW_SEGMENT_CHARS} 字/块）"
                )
            sheet_segments = build_table_segments_from_rows(
                row_matrix,
                table_index,
                offset_ref,
                parts,
                title,
                table_body_html=table_html,
                row_batch_size=row_batch,
            )
            for seg in sheet_segments:
                meta = dict(seg.metadata or {})
                meta["sheet"] = title
                meta["excel_engine"] = "html_table"
                seg.metadata = meta
            segments.extend(sheet_segments)
            table_index += 1

    full_text = "\n".join(parts).strip()
    if not full_text and segments:
        full_text = "\n".join(s.text for s in segments)
    parser_type = suffix if suffix in {"xlsx", "xlsm", "xls"} else "xlsx"
    _log(f"Excel HTML 解析汇总：分段 {len(segments)}，文本长度 {len(full_text)}")
    return ParsedDocument(
        parser_type=parser_type,
        text=full_text,
        char_count=len(full_text),
        segments=segments or [ParsedSegment(text="", start_offset=0, end_offset=0, metadata={})],
        metadata={"sheet_count": table_index, "excel_engine": "html_table", "table_count": table_index, "chunking_strategy": "table_segments"},
    )


def _parse_excel(
    file_bytes: bytes,
    suffix: str,
    *,
    engine: str = "html_table",
    log: Callable[[str], None] | None = None,
) -> ParsedDocument:
    if engine == "tsv":
        return _parse_excel_tsv(file_bytes, suffix, log=log)
    return _parse_excel_html_table(file_bytes, suffix, log=log)


def _parse_csv(file_bytes: bytes) -> ParsedDocument:
    text = _decode_text_content(file_bytes)
    sio = StringIO(text)
    rows: list[str] = []
    try:
        sample = text[:4096]
        dialect = csv.Sniffer().sniff(sample) if sample.strip() else csv.excel
    except csv.Error:
        dialect = csv.excel
    sio.seek(0)
    for row in csv.reader(sio, dialect):
        if any((c or "").strip() for c in row):
            rows.append("\t".join((c or "").strip() for c in row))
    full_text = "\n".join(rows).strip()
    segment = ParsedSegment(text=full_text, start_offset=0, end_offset=len(full_text), metadata={"block": "csv"})
    return ParsedDocument(
        parser_type="csv",
        text=full_text,
        char_count=len(full_text),
        segments=[segment],
        metadata={"row_count": len(rows)},
    )


def _metadata_log_preview(metadata: dict | None, *, max_len: int = 800) -> str:
    """日志用 metadata 摘要，排除 MinerU 侧车大字段避免 SSE 刷屏。"""
    if not metadata:
        return "{}"
    slim = {k: v for k, v in metadata.items() if not str(k).startswith(("_mineru_", "_parse_"))}
    preview = json.dumps(slim, ensure_ascii=False)
    if len(preview) > max_len:
        preview = preview[:max_len] + "…"
    return preview


def _parse_html_word_doc(file_bytes: bytes, *, file_name: str, log: Callable[[str], None] | None = None) -> ParsedDocument:
    """Parse Word-exported HTML saved with a .doc extension (fallback when LO convert fails)."""

    def _log(msg: str) -> None:
        if log:
            log(msg)

    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self._parts: list[str] = []
            self._skip_depth = 0

        def handle_starttag(self, tag: str, attrs) -> None:
            if tag.lower() in {"script", "style"}:
                self._skip_depth += 1

        def handle_endtag(self, tag: str) -> None:
            if tag.lower() in {"script", "style"} and self._skip_depth:
                self._skip_depth -= 1

        def handle_data(self, data: str) -> None:
            if self._skip_depth:
                return
            text = html_module.unescape(data.strip())
            if text:
                self._parts.append(text)

    charset = "utf-8"
    head_sample = file_bytes[:4096].decode("latin-1", errors="ignore").lower()
    charset_match = re.search(r'charset\s*=\s*["\']?([\w-]+)', head_sample)
    if charset_match:
        charset = charset_match.group(1)
    try:
        html_text = file_bytes.decode(charset)
    except (LookupError, UnicodeDecodeError):
        for encoding in ("gb18030", "gbk", "utf-8", "latin-1"):
            try:
                html_text = file_bytes.decode(encoding)
                break
            except UnicodeDecodeError:
                html_text = ""
        else:
            html_text = file_bytes.decode("utf-8", errors="ignore")

    parser = _TextExtractor()
    parser.feed(html_text)
    body_text = "\n".join(parser._parts).strip()
    if not body_text:
        raise UnsupportedDocumentError(
            "该 .doc 为 HTML 网页格式，但未能提取正文。请用 Word/WPS 打开后另存为 .docx 再上传。"
        )

    _log(f"HTML 型 .doc：提取正文 {len(body_text)} 字符")
    segments: list[ParsedSegment] = []
    offset = 0
    for idx, block in enumerate(part for part in body_text.splitlines() if part.strip()):
        start = offset
        end = start + len(block)
        segments.append(
            ParsedSegment(
                text=block,
                start_offset=start,
                end_offset=end,
                metadata={"block_index": idx, "source_format": "doc_html"},
            )
        )
        offset = end + 1

    return ParsedDocument(
        parser_type="doc",
        text=body_text,
        char_count=len(body_text),
        segments=segments,
        metadata={
            "source_format": "doc",
            "converted_via": "html_text_extract",
            "parser_backend": "native",
            "original_file_name": file_name,
        },
    )


def parse_document(
    file_name: str,
    file_bytes: bytes,
    log: Callable[[str], None] | None = None,
    *,
    pdf_parser: str | None = None,
    pdf_parser_hybrid: bool | None = None,
    parser_options: Any | None = None,
) -> ParsedDocument:
    from app.core.parser_config import ParserOptions, resolve_parsers

    def _log(msg: str) -> None:
        if log:
            log(msg)

    opts: ParserOptions | None = parser_options if isinstance(parser_options, ParserOptions) else None
    if opts is None and (pdf_parser is not None or pdf_parser_hybrid is not None):
        legacy_cfg: dict[str, Any] = {}
        if pdf_parser is not None:
            legacy_cfg["pdf_parser"] = pdf_parser
        if pdf_parser_hybrid is not None:
            legacy_cfg["pdf_parser_hybrid"] = pdf_parser_hybrid
        opts = resolve_parsers(legacy_cfg)
    elif opts is None:
        opts = resolve_parsers(None)

    suffix = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
    _log(f"开始解析「{file_name}」，扩展名 .{suffix or '（无）'}，大小 {len(file_bytes)} 字节")

    if suffix == "doc":
        from app.core.doc_conversion import convert_doc_to_docx_bytes, detect_doc_container

        container = detect_doc_container(file_bytes)
        if container == "html":
            _log("检测到 HTML 型 .doc（Word 网页导出），正在转换为 docx…")
        else:
            _log("检测到旧版 Word（.doc），正在转换为 .docx…")
        try:
            docx_bytes = convert_doc_to_docx_bytes(file_bytes, source_name=file_name, log=_log)
        except UnsupportedDocumentError:
            if container == "html":
                _log("LibreOffice 转换失败，改用 HTML 正文提取…")
                return _parse_html_word_doc(file_bytes, file_name=file_name, log=log)
            raise
        _log(f".doc 已转换，继续按 docx 解析（引擎={opts.docx_engine}）…")
        if opts.docx_engine == "mineru":
            docx_name = file_name.rsplit(".", 1)[0] + ".docx" if "." in file_name else f"{file_name}.docx"
            doc = _parse_with_mineru(
                docx_bytes,
                suffix="docx",
                file_name=docx_name,
                parser_type="docx",
                log=log,
                require_format_whitelist=False,
            )
        else:
            _log(f"使用 {opts.docx_engine} 解析 docx…")
            doc = _parse_docx(docx_bytes)
            _log(f"docx：字符数 {doc.char_count}，分段 {len(doc.segments)}")
        meta = dict(doc.metadata or {})
        meta["source_format"] = "doc"
        if meta.get("converted_via") is None:
            meta["converted_via"] = "libreoffice"
        if opts.docx_engine == "mineru":
            meta["parser_backend"] = "mineru"
        elif not meta.get("parser_backend"):
            meta["parser_backend"] = opts.docx_engine or "python-docx"
        doc.metadata = meta
        doc.parser_type = "doc"
        return doc

    if suffix == "docx":
        if opts.docx_engine == "mineru":
            return _parse_with_mineru(
                file_bytes,
                suffix=suffix,
                file_name=file_name,
                parser_type="docx",
                log=log,
                require_format_whitelist=False,
            )
        _log(f"使用 {opts.docx_engine} 解析 docx…")
        doc = _parse_docx(file_bytes)
        _log(f"docx：字符数 {doc.char_count}，分段 {len(doc.segments)}")
        return doc

    if suffix in {"xlsx", "xlsm", "xls"}:
        excel_suffix = suffix
        if opts.excel_engine == "mineru":
            return _parse_with_mineru(
                file_bytes,
                suffix=excel_suffix,
                file_name=file_name,
                parser_type=excel_suffix,
                log=log,
                require_format_whitelist=False,
            )
        _log(f"Excel 解析引擎：{opts.excel_engine}")
        doc = _parse_excel(file_bytes, excel_suffix, engine=opts.excel_engine, log=log)
        _log(f"Excel：字符数 {doc.char_count}，分段 {len(doc.segments)}")
        return doc

    if suffix == "csv":
        if opts.csv_engine == "mineru":
            return _parse_with_mineru(
                file_bytes,
                suffix=suffix,
                file_name=file_name,
                parser_type="csv",
                log=log,
                require_format_whitelist=False,
            )
        _log("按 CSV 规则解析…")
        doc = _parse_csv(file_bytes)
        _log(f"CSV：有效数据行 {doc.metadata.get('row_count', 0)}，文本长度 {doc.char_count}")
        return doc

    if suffix == "pdf":
        doc = _parse_pdf(
            file_bytes,
            file_name=file_name,
            pdf_parser=opts.pdf_engine,
            pdf_parser_hybrid=opts.pdf_hybrid if opts.pdf_engine == "opendataloader" else None,
            log=log,
        )
        meta_for_log = {
            k: v for k, v in (doc.metadata or {}).items() if not str(k).startswith(("_mineru_", "_parse_"))
        }
        meta_preview = json.dumps(meta_for_log, ensure_ascii=False)
        if len(meta_preview) > 800:
            meta_preview = meta_preview[:800] + "…"
        _log(f"PDF：字符数 {doc.char_count}，分段 {len(doc.segments)}，元数据 {meta_preview}")
        return doc

    if suffix in _IMAGE_SUFFIXES:
        return _parse_image(file_bytes, suffix=suffix, file_name=file_name, log=log)

    text = _decode_text_content(file_bytes).strip()
    if not text:
        _log("解码后文本为空")
        return ParsedDocument(parser_type=suffix or "txt", text="", char_count=0, segments=[], metadata={})

    if suffix == "md":
        if opts.text_engine == "mineru":
            return _parse_with_mineru(
                file_bytes,
                suffix="md",
                file_name=file_name,
                parser_type="md",
                log=log,
                require_format_whitelist=False,
            )
        segments = _build_segments_from_sections(_split_markdown_sections(text))
        _log(f"Markdown：分段数 {len(segments)}")
        return ParsedDocument(
            parser_type="md",
            text=text,
            char_count=len(text),
            segments=segments,
            metadata={"section_count": len(segments)},
        )

    if suffix == "txt" and opts.text_engine == "mineru":
        return _parse_with_mineru(
            file_bytes,
            suffix="txt",
            file_name=file_name,
            parser_type="txt",
            log=log,
            require_format_whitelist=False,
        )

    segments = [
        ParsedSegment(
            text=text,
            start_offset=0,
            end_offset=len(text),
            metadata={},
        )
    ]
    _log(f"纯文本：长度 {len(text)}")
    return ParsedDocument(parser_type=suffix or "txt", text=text, char_count=len(text), segments=segments, metadata={})
