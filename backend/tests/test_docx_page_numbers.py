"""Word 文档真实页码（分页符 / lastRenderedPageBreak）测试。"""

from io import BytesIO

from docx import Document

from app.core.document_parser import _parse_docx


def test_hard_page_break_assigns_page_numbers():
    doc = Document()
    doc.add_paragraph("Page one")
    doc.add_page_break()
    doc.add_paragraph("Page two")
    buf = BytesIO()
    doc.save(buf)

    parsed = _parse_docx(buf.getvalue())
    pages = [s.page_no for s in parsed.segments if s.text.strip()]
    assert pages == [1, 2]
